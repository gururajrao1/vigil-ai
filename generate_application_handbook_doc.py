#!/usr/bin/env python3
"""Generate VigilAI Application Handbook as a Word document.

Usage:
    pip install python-docx
    python generate_application_handbook_doc.py

Output:
    docs/VigilAI_Application_Handbook.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x0F, 0x2C, 0x59)
CHARCOAL = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CALLOUT = "E8EEF6"
TIP = "EEF6F0"

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "docs" / "VigilAI_Application_Handbook.docx"
MD = ROOT / "docs" / "VIGILAI_APPLICATION_HANDBOOK.md"


def _font(run, *, size=11, bold=False, color=CHARCOAL, name="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rpr.append(rfonts)


def _shade(cell, hex_fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _p(doc, text, *, size=11, bold=False, color=CHARCOAL, space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    _font(run, size=size, bold=bold, color=color)
    return p


def _h(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _font(run, size=sizes.get(level, 12), bold=True, color=NAVY)
    return p


def _callout(doc, title: str, body: str, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _shade(cell, fill)
    cell.text = ""
    p0 = cell.paragraphs[0]
    r0 = p0.add_run(title)
    _font(r0, size=10, bold=True, color=NAVY)
    p1 = cell.add_paragraph()
    r1 = p1.add_run(body)
    _font(r1, size=10, color=CHARCOAL)
    doc.add_paragraph()


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _shade(cell, "0F2C59")
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _font(run, size=9, bold=True, color=WHITE)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _font(run, size=9, color=CHARCOAL)
    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    _p(doc, "VigilAI", size=28, bold=True, color=NAVY, space_after=2,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "Application Handbook", size=16, bold=True, color=NAVY, space_after=8,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(
        doc,
        "Where · How · What · Why — features, metrics, data sources, and empty-result teaching scripts",
        size=11, color=MUTED, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _p(
        doc,
        "Live: https://vigil-ai-eight.vercel.app  ·  API: https://vigil-ai-api.onrender.com",
        size=9, color=MUTED, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _callout(
        doc,
        "Primary markdown (diagrams + §19 empty-result scripts)",
        "Full mermaid architecture diagrams, metrics explanations, data-source cheat sheet, "
        "and “when you don’t see expected output” live in docs/VIGILAI_APPLICATION_HANDBOOK.md. "
        "This Word file is the shareable handout for stakeholders who prefer .docx.",
        TIP,
    )

    # --- 1 ---
    _h(doc, "1. What VigilAI is")
    _p(
        doc,
        "VigilAI is an offline-first pharmacovigilance and device-vigilance platform. "
        "It ingests unstructured patient and regulatory chatter, extracts clinical entities, "
        "validates adverse events with an explainable 4-gate engine, computes disproportionality "
        "(PRR / ROR / χ² / EBGM / BCPNN), overlays analytic lenses (Remine, DDI, Pregnancy, "
        "Risk populations…), and manages signals through a GVP-shaped workflow with demo "
        "E2B / CIOMS / SAR exports.",
    )
    _table(
        doc,
        ["It is", "It is not"],
        [
            ["Working prototype for demos & KT", "Validated regulatory submission system"],
            ["Drugs + vaccines + devices", "US-only or drug-only"],
            ["Offline-capable (every network path degrades)", "Dependent on paid API keys"],
            ["Honest about licensed-network surrogates", "A live WHO VigiBase bulk feed"],
        ],
    )

    # --- 2 ---
    _h(doc, "2. Why it exists")
    _p(
        doc,
        "Pre-market trials are small and short. Rare and late harms often appear first in "
        "conversation before structured ICSRs. VigilAI hears that chatter, codes it to "
        "MedDRA-style / ATC / GMDN concepts, decides “is this an AE?” with auditable gates, "
        "ranks associations with regulator-shaped stats, stress-tests them (remine, DDI, "
        "pregnancy, risk strata), and hands off to workflow + demo exports.",
    )

    # --- 3 ---
    _h(doc, "3. Roles & login")
    _table(
        doc,
        ["Role", "Account", "Can do"],
        [
            ["Admin", "admin@vigilai.dev / admin123", "Everything (users, ingest, reset)"],
            ["Analyst", "analyst@vigilai.dev / analyst123", "Analytics, review, Forge, Projects, Pathfinder"],
            ["Viewer", "viewer@vigilai.dev / viewer123", "Read signals / lenses / exports"],
        ],
    )

    # --- 4 ---
    _h(doc, "4. Architecture (summary)")
    _p(doc, "Layers", size=12, bold=True, color=NAVY)
    _p(
        doc,
        "Frontend (React + Vite hubs, ⌘K) → FastAPI + JWT → NLP / 4-gate AE → "
        "Disproportionality + analytic lenses → Lifecycle / alerts → SQLite or Neon Postgres. "
        "Optional openFDA / PubMed / DailyMed / MAUDE degrade to fixtures offline. "
        "Deploy: Vercel (UI) → Render (API) → Neon (DB).",
    )
    _p(doc, "End-to-end pipeline", size=12, bold=True, color=NAVY)
    _p(
        doc,
        "Ingest → PII scrub → entities / sentiment / negation → 4-gate AE → recompute signals "
        "(PRR/ROR/χ²/EBGM/IC) → lenses (read-only overlays) → background evidence enrich → "
        "workflow → SAR / E2B / CIOMS.",
    )
    _p(doc, "Remine (competition bias)", size=12, bold=True, color=NAVY)
    _p(
        doc,
        "Case-level unmasking drops whole reports mentioning a masker product, then rebuilds "
        "the 2×2. Masking ratio MR = PRR_after/PRR_before = coreporting_term × comparator_term. "
        "Only threshold-crossing (unmasked) is strongly actionable; raw PRR rises often reflect "
        "shared comparator arithmetic.",
    )

    # --- 5 ---
    _h(doc, "5. How to use (first 5 minutes)")
    for i, step in enumerate(
        [
            "Open the live app → Login (admin or analyst).",
            "Confirm project in the header (General Pharmacovigilance by default).",
            "Data Sources → Load PV demo pack if Remine / DDI / Register look thin.",
            "Safety Signals → Detect → type warfarin or paracetamol → open a row.",
            "Read Conclusions first, then the label line, then PrOACT, then DMA.",
            "Safety Signals → Register → advance one lifecycle action (Looking into it).",
            "Dashboard → Inspection & COU → overdue SLAs + Not validated for boundaries.",
            "Lenses → Risk → REM: try paracetamol + Hepatic injury (ranked) and optionally pacemaker (empty gates).",
            "Optional: Lenses → Remine lab → Needs review → Run remine on warfarin / haemorrhage.",
        ],
        1,
    ):
        _p(doc, f"{i}. {step}", space_after=3)

    # --- 6 ---
    _h(doc, "6. Navigation map")
    _table(
        doc,
        ["Hub", "Route", "Purpose"],
        [
            ["Dashboard", "/dashboard", "Corpus · Ops KPIs · Inspection & COU"],
            ["Safety Signals", "/signals", "Detect · Register · Workflow · Alerts"],
            ["Analytic Lenses", "/lenses", "Predictive · Remine · REM Risk · DDI · Pregnancy · SMQ · Class · Vaccine · Geo · vs FAERS"],
            ["Evidence Explorer", "/graph", "Drug↔AE graph · Story · Glossary"],
            ["Projects", "/projects", "Workspaces + keywords for Pathfinder"],
            ["Source Discovery", "/source-queue", "Pathfinder · Manual forum URL"],
            ["Data Sources", "/sources", "Catalog · Live stream · Networks · Agent"],
            ["Data Forge", "/forge", "Synthetic posts (analyst+)"],
            ["Users", "/users", "Admin only"],
        ],
    )

    # --- how to use each feature ---
    _h(doc, "7. How to use each feature (step by step)")
    _p(
        doc,
        "Full click-by-click recipes live in the markdown handbook §7 "
        "(Dashboard, Detect, Register, Signal Detail, Workflow, Alerts, Remine, REM Risk, "
        "Predictive intel, DDI, Pregnancy, SMQ, Class, Vaccine, Geo, vs FAERS, Graph, "
        "Projects/keywords, Pathfinder, Sources, Forge, Users). Summary:",
    )
    for line in [
        "Detect: search product → open row → read Conclusions + label first.",
        "Register: GVP IX queue — plain lifecycle buttons (Inbox → Looking into it → …).",
        "Detail: Conclusions · PrOACT · Inspection · Triangulation · Causality · always-on PGx.",
        "REM ranking: paracetamol + Hepatic injury shows strata; pacemaker often correctly empty.",
        "Remine lab: filter Needs review → Run remine → judge threshold crossing, not raw PRR rise.",
        "Inspection & COU: SLA overdue + hard Not validated for boundaries (not a backlog).",
        "Sources: Load PV demo pack once when Remine/DDI/REM look thin.",
        "Exports: SAR / PBRER / E2B / CIOMS — demo templates only.",
    ]:
        _p(doc, "• " + line, space_after=2)

    _h(doc, "8. Feature catalog (highlights)")
    _table(
        doc,
        ["Feature", "What / Why"],
        [
            ["Detect + jump search", "Ranked product→event table; type to jump (no scrolling)"],
            ["Conclusions / briefing", "Plain-English verdict from PRR, n, calibration"],
            ["Label filter", "Novel vs in-label vs boxed; Weber gate on launch noise"],
            ["Register", "GVP IX tracking queue + SAR/PBRER exports"],
            ["4-gate AE", "Product · symptom · negative sentiment · non-negated"],
            ["Remine lab", "Corpus-wide competition-bias screen; read-only"],
            ["REM ranking", "Subpopulation Risk Elevation Multiplier (gates REM≥1.5, χ²≥4)"],
            ["PGx", "Always-on CPIC/PharmGKB screen (actionable or clean)"],
            ["PrOACT / BRAT", "Benefit vs severe AE balance + eight PrOACT-URL dimensions"],
            ["Inspection / COU", "SLA overdue + Context of Use credibility scorecard"],
            ["Lot / longitudinal", "Shown only when lot cues or biologic/ATMP/late-signal"],
            ["DDI / Pregnancy", "Polypharmacy co-mentions · teratogen / congenital cohort"],
            ["SMQ / Class / Vaccine / Geo / vs FAERS", "Syndrome, ATC class, AESI, clusters, FDA divergence"],
            ["Project keywords", "Intent vocabulary for Pathfinder + literature narrowing"],
        ],
    )

    # --- 8 ---
    _h(doc, "9. Signal science (cheat sheet)")
    _table(
        doc,
        ["Metric", "Meaning"],
        [
            ["PRR / ROR", "Reporting rate / odds vs rest of corpus (+0.5 continuity)"],
            ["χ² (Yates)", "Independence on 2×2; ≥4 supports signal"],
            ["EB05 / IC025", "Bayesian lower bounds (MGPS / BCPNN flavored)"],
            ["SDR", "Composite signal of disproportionate reporting"],
            ["STRONG", "PRR≥2, χ²≥4, count≥3"],
            ["REM", "P(AE|Drug∩Subpop) / P(AE|Drug∩General); keep if REM≥1.5 & χ²≥4"],
            ["Remine outcomes", "unmasked · co_reported · vanished · attenuated · amplified · stable"],
            ["NNT / NNH", "Illustrative benefit / report-derived harm proxy in PrOACT"],
        ],
    )

    _h(doc, "9b. When output looks empty (not always a bug)")
    _p(
        doc,
        "Full teaching scripts: markdown handbook §19. Short version for demos:",
    )
    for line in [
        "REM empty + n_exposed shown: gates held (tiny n / high baseline) — try paracetamol → Hepatic injury.",
        "Lot / longitudinal missing: correctly not relevant (no lot cues / not biologic-ATMP).",
        "PGx “no Level-A match”: clean screen (e.g. apixaban → haemorrhage).",
        "Triangulation INSUFFICIENT: social loud, regulatory thin — teaching point.",
        "Red error / endless spinner / 404 on /api/inspection: deploy or wake API — real failure.",
        "Detect Novel chip ≠ PrOACT — open Signal Detail for the PrOACT-URL / BRAT badge.",
    ]:
        _p(doc, "• " + line, space_after=2)

    # --- 9 ---
    _h(doc, "10. Project keywords — why they matter")
    _p(
        doc,
        "Three layers: (A) Projects → keywords drive Pathfinder + literature retrieval; "
        "(B) the same product/event strings type as-is into Detect, Remine, Risk, Graph; "
        "(C) ontology resolves brand↔generic↔chemical on Signal Detail / API. "
        "Use 3–8 project terms mixing product/class, event, and community language. After saving, "
        "activate the project → Source Discovery → Run Pathfinder → Approve → Fetch.",
    )

    # --- 10 ---
    _h(doc, "11. Compiled keyword packs (copy/paste)")
    packs = [
        ("General PV", "adverse reaction, side effect, drug safety, pharmacovigilance, patient forum, MedWatch"),
        ("Anticoagulants", "warfarin, rivaroxaban, apixaban, haemorrhage, bleeding, anticoagulant, INR"),
        ("Psychiatry", "paroxetine, sertraline, lithium, suicidal ideation, akathisia, SSRI, bipolar forum"),
        ("GLP-1 / metabolic", "semaglutide, Ozempic, Wegovy, pancreatitis, gastroparesis, nausea, weight loss drug"),
        ("Oncology / ICI", "pembrolizumab, nivolumab, checkpoint inhibitor, immune-related AE, colitis, pneumonitis, oncology forum"),
        ("Pregnancy", "pregnancy, congenital anomaly, birth defect, teratogen, lithium pregnancy, valproate, neural tube defect"),
        ("Vaccine / AESI", "myocarditis, vaccine side effects, reactogenicity, MMR, COVID-19 vaccine, VAERS, immunization"),
        ("Devices · cardiac", "pacemaker, coronary stent, defibrillator, lead fracture, device malfunction, MAUDE, implant forum"),
        ("Devices · diabetes", "insulin pump, continuous glucose monitor, CGM, overinfusion, sensor error, infusion set, diabetes device"),
        ("DDI / polypharmacy", "polypharmacy, drug interaction, warfarin amiodarone, serotonin syndrome, concomitant medication"),
    ]
    _table(doc, ["Pack", "Keywords"], packs)

    _p(doc, "In-app Detect jump strings (corpus lookup, not Pathfinder)", size=12, bold=True, color=NAVY)
    _p(
        doc,
        "Drugs: warfarin, rivaroxaban, semaglutide, paroxetine, sertraline, lithium, ibuprofen.  "
        "Vaccines: covid-19 mrna vaccine, MMR.  "
        "Devices: coronary stent, catheter, pacemaker, insulin pump, continuous glucose monitor.  "
        "Events: Haemorrhage, Fatigue, Anxiety, Device malfunction, Nausea.",
    )

    # --- 11 ---
    _h(doc, "12. Quick keyword index (find anything)")
    _table(
        doc,
        ["You want…", "Search / say", "Go to"],
        [
            ["Main signal table", "Detect, SDR, PRR", "/signals"],
            ["GVP Register", "Register, PBRER", "/signals?tab=register"],
            ["Conclusions / PrOACT", "Conclusions, PrOACT", "Signal Detail"],
            ["Label novel vs in-label", "label filter, Novel", "Signal Detail hero"],
            ["PGx", "CPIC, PharmGKB", "Signal Detail"],
            ["Inspection / COU", "SLA, credibility", "/dashboard?tab=governance"],
            ["REM ranking", "REM, risk populations", "/lenses?tab=risk"],
            ["Competition bias", "Remine, masking", "/lenses?tab=remine"],
            ["Empty REM explanation", "§19 empty result", "Handbook §19"],
            ["Demo data", "PV demo pack", "/sources"],
            ["Export ICSR", "E2B, CIOMS, SAR, PBRER", "Detail / Register"],
        ],
    )

    # --- 12 ---
    _h(doc, "13. Disclaimers")
    _p(
        doc,
        "Prototype; synthetic data is fictional; openFDA = US FAERS/MAUDE (plus other open feeds "
        "as wired); MedDRA coding is an open surrogate (not licensed MedDRA); E2B/CIOMS/SAR/PBRER "
        "are demo templates; PGx and PrOACT NNT/NNH are teaching surrogates; COU “Not validated for” "
        "items are permanent scope boundaries; not for clinical use.",
    )

    _p(
        doc,
        "Full diagrams, metrics, data-source cheat sheet & empty-result scripts: docs/VIGILAI_APPLICATION_HANDBOOK.md",
        size=9, color=MUTED, space_after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
