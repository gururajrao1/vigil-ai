"""One-click DOCX / Markdown export for regulatory drafts."""
from __future__ import annotations

import io
from typing import Optional, Tuple


def render_docx_or_markdown(
    title: str,
    markdown_body: str,
    *,
    prefer_docx: bool = True,
) -> Tuple[bytes, str, str]:
    """Return (bytes, media_type, filename_ext).

    Tries python-docx when available; otherwise returns UTF-8 markdown.
    """
    if prefer_docx:
        try:
            from docx import Document  # type: ignore

            doc = Document()
            doc.add_heading(title, 0)
            for block in (markdown_body or "").split("\n"):
                if block.startswith("# "):
                    doc.add_heading(block[2:].strip(), level=1)
                elif block.startswith("## "):
                    doc.add_heading(block[3:].strip(), level=2)
                elif block.startswith("> "):
                    doc.add_paragraph(block[2:].strip())
                elif block.strip():
                    doc.add_paragraph(block)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue(), (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ), "docx"
        except Exception:
            pass
    return (
        (markdown_body or "").encode("utf-8"),
        "text/markdown; charset=utf-8",
        "md",
    )
